import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import tempfile, os, re
from pathlib import Path
from subprocess import run, PIPE
from PyPDF2 import PdfMerger

st.set_page_config(page_title="📄→📦 UniUni BOL", layout="wide")
st.title("UniUni BOL Generator (DOCX→PDF via LibreOffice)")


# 替换 Word 文档里的占位符
def replace_all_text(doc: Document, reps: dict):
    for para in doc.paragraphs:
        runs = para.runs
        txt = "".join(r.text for r in runs)
        for k, v in reps.items():
            if k in txt:
                txt = txt.replace(k, v)
                for r in runs:
                    r.text = ""
                runs[0].text = txt
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    runs = para.runs
                    txt = "".join(r.text for r in runs)
                    for k, v in reps.items():
                        if k in txt:
                            txt = txt.replace(k, v)
                            for r in runs:
                                r.text = ""
                            runs[0].text = txt


# 调用 LibreOffice headless 转 PDF
def convert_doc_to_pdf_native(doc_file: Path, outdir: Path, timeout: int = 60):
    try:
        proc = run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(outdir),
                str(doc_file),
            ],
            stdout=PIPE,
            stderr=PIPE,
            timeout=timeout,
            check=True,
        )
        out = proc.stdout.decode("utf-8")
        m = re.search(r"-> (.*?) using filter", out)
        if m:
            return Path(m.group(1)), None
        return None, "LibreOffice 未返回输出文件路径"
    except Exception as e:
        return None, e


# UI 输入
ship_date = st.date_input("Ship Date")
uploaded_xl = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])
uploaded_docx = st.file_uploader("Upload Word Template (.docx)", type=["docx"])

if ship_date and uploaded_xl and uploaded_docx:
    if st.button("🚀 Generate Combined PDF"):
        # 读 Excel
        df = pd.read_excel(BytesIO(uploaded_xl.read()))
        needed = {"Address", "Phone", "Note", "DSP"}
        missing = needed - set(df.columns)
        if missing:
            st.error(f"❌ Excel 缺少列: {sorted(missing)}")
        else:
            total = len(df)
            status = st.empty()  # 占位符：用于覆盖上一次的进度消息
            prog = st.progress(0)  # 进度条
            pdfs = []
            with tempfile.TemporaryDirectory() as tmp:
                tpl_path = Path(tmp) / "template.docx"
                tpl_path.write_bytes(uploaded_docx.read())

                for i, row in df.iterrows():
                    seq = i + 1
                    doc = Document(tpl_path)

                    reps = {
                        "SEA-[pickup address]+TEPHONE+NOTE": f"SEA - {row['Address']} | TEL: {row['Phone']} | Note: {row['Note']}",
                        "UNI-SEA-PICKUP-MM/DD/YYYY-SEQ": f"UNI-SEA-PICKUP-{ship_date.strftime('%m/%d/%Y')}-{seq}",
                        "Carrier Name: GN GREENWHEELS INC.": f"Carrier Name: GN GREENWHEELS INC. - {row['DSP']}",
                        "Ship_date": ship_date.strftime("%m/%d/%Y"),
                    }
                    replace_all_text(doc, reps)

                    docx_out = Path(tmp) / f"{seq}.docx"
                    doc.save(docx_out)

                    out_pdf, err = convert_doc_to_pdf_native(docx_out, Path(tmp))
                    if err or not out_pdf or not out_pdf.exists():
                        st.error(f"❌ 第 {seq} 行转换失败: {err}")
                        break

                    pdfs.append(str(out_pdf))
                    # 覆盖前一次消息（1/20 → 2/20…）
                    status.success(f"✅ {seq} / {total} done")
                    # 更新进度条
                    prog.progress(int(seq / total * 100))

                else:
                    # 全部成功后合并 PDF
                    merger = PdfMerger()
                    for p in pdfs:
                        merger.append(p)
                    buf = BytesIO()
                    merger.write(buf)
                    merger.close()
                    buf.seek(0)

                    status.empty()  # 清掉最后一条进度消息
                    prog.empty()  # 隐藏进度条
                    st.success("✅ All done! Download below.")
                    st.download_button(
                        "Download All BOLs",
                        data=buf,
                        file_name="All_BOLs.pdf",
                        mime="application/pdf",
                    )
